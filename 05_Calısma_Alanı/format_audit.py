from docx import Document
import re

doc = Document(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx")
issues = []

# 1. TITLE
title = doc.paragraphs[0].text.strip()
print(f'1. BASLIK: "{title}"')
print(f'   Kelime: {len(title.split())} (sablon max 11)')
if len(title.split()) > 11:
    issues.append(f'BASLIK: {len(title.split())} kelime > 11')
print()

# 2. FONT
font_counts = {}
for p in doc.paragraphs:
    for run in p.runs:
        f = run.font.name or 'Inherited'
        font_counts[f] = font_counts.get(f, 0) + len(run.text)
print('2. YAZI TIPI:')
for f, c in sorted(font_counts.items(), key=lambda x: -x[1])[:5]:
    print(f'   {f}: {c} karakter')
print()

# 3. SECTION HEADINGS
print('3. BOLUM BASLIKLARI (sablon: ana=koyu, alt=italik):')
section_re = re.compile(r'^(\d+\.[\d.]*)\s+(.+)')
for p in doc.paragraphs:
    m = section_re.match(p.text.strip())
    if m:
        num, ttl = m.group(1), m.group(2)
        bold_runs = [r.bold for r in p.runs if r.text.strip()]
        italic_runs = [r.italic for r in p.runs if r.text.strip()]
        is_bold = all(bold_runs) if bold_runs else False
        is_italic = all(italic_runs) if italic_runs else False
        dots = num.count('.')
        expected = 'KOYU' if dots == 1 else 'ITALIK'
        actual = 'KOYU' if is_bold else ('ITALIK' if is_italic else 'NORMAL')
        status = 'OK' if expected == actual else 'HATA'
        if status == 'HATA':
            issues.append(f'{num}: beklenen={expected}, gercek={actual}')
        print(f'   {num} [{actual}] beklenen:{expected} {status} - {ttl[:40]}')
print()

# 4. ABSTRACT ITALIC
print('4. ABSTRACT ITALIK:')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Abstract':
        if i+1 < len(doc.paragraphs):
            nx = doc.paragraphs[i+1]
            italic_runs = [r.italic for r in nx.runs if r.text.strip()]
            is_it = all(italic_runs) if italic_runs else False
            print(f'   Ingilizce ozet italik: {is_it}')
            if not is_it:
                issues.append('Ingilizce Abstract metni italik olmali')
        break
print()

# 5. KEYWORDS ALPHABETICAL
print('5. ANAHTAR SOZCUKLER SIRALAMA:')
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('Anahtar'):
        kws = t.replace('Anahtar sozcukler:', '').replace('Anahtar Sozcukler:', '').strip()
        kws = t.split(':',1)[1].strip().rstrip('.').split(',')
        kws = [k.strip() for k in kws]
        tr_key = lambda x: x.lower().replace('c','c1').replace('g','g1').replace('i','i1').replace('o','o1').replace('s','s1').replace('u','u1')
        srt = sorted(kws, key=tr_key)
        if kws == srt:
            print(f'   Alfabetik OK')
        else:
            print(f'   SIRA HATASI')
            print(f'   Mevcut:  {kws}')
            print(f'   Olmasi:  {srt}')
            issues.append('TR anahtar sozcukler alfabetik degil')
    if t.startswith('Keywords:'):
        kws = t.split(':',1)[1].strip().rstrip('.').split(',')
        kws = [k.strip() for k in kws]
        srt = sorted(kws, key=str.lower)
        if kws == srt:
            print(f'   EN Keywords alphabetical OK')
        else:
            print(f'   EN Keywords NOT alphabetical')
            print(f'   Current: {kws}')
            print(f'   Should:  {srt}')
            issues.append('EN keywords not alphabetical')
print()

# 6. FIGURE CAPTIONS - first word only capitalized per template
print('6. SEKIL/CIZELGE BASLIK KURALLARI:')
cap_issues = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if re.match(r'^(Sekil|Cizelge)\s+\d+\.', t, re.IGNORECASE):
        # Check if bold
        bold_runs = [r.bold for r in p.runs if r.text.strip()]
        is_bold = any(bold_runs) if bold_runs else False
        if not is_bold and cap_issues < 5:
            print(f'   KOYU DEGIL: {t[:80]}')
            cap_issues += 1
if cap_issues == 0:
    print('   Tum basliklar koyu OK')
print()

# 7. PAGE MARGINS and other
print('7. DIGER KONTROLLER:')
# Check if formulas are left-aligned per template
formula_re = re.compile(r'\(\d+\)$')
for p in doc.paragraphs:
    if formula_re.search(p.text.strip()):
        print(f'   Formula: {p.text.strip()[:80]}')
        break

# Check word count
total = sum(len(p.text.split()) for p in doc.paragraphs)
print(f'   Kelime: {total}')

print()
print('='*60)
print(f'TOPLAM SORUN: {len(issues)}')
print('='*60)
for iss in issues:
    print(f'  > {iss}')
