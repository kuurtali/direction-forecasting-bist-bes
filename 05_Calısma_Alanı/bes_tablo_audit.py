# -*- coding: utf-8 -*-
"""
BES_Pension_Fund_Report.docx — Tablo B2 ve B3 DERIN AUDIT
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from pathlib import Path

BASE = Path(r"c:\Users\Kurt\Desktop\Proje")
bes_path = BASE / "01_Savunma_ve_Ana_Metinler" / "BES_Pension_Fund_Report.docx"

doc = Document(str(bes_path))

print("=" * 80)
print("BES_Pension_Fund_Report.docx — TABLO TARAMASI")
print("=" * 80)
print(f"Toplam tablo sayisi: {len(doc.tables)}")
print(f"Toplam paragraf sayisi: {len(doc.paragraphs)}")

# Tum tablolari listele
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    # Ilk satirin icerigini al (baslik)
    header = []
    if rows > 0:
        for cell in table.rows[0].cells:
            header.append(cell.text.strip()[:30])
    print(f"\n--- TABLO {i+1} ({rows} satir x {cols} sutun) ---")
    print(f"  Baslik: {' | '.join(header)}")
    
    # Eger AMZ iceriyorsa detayli goster
    has_amz = False
    for row in table.rows:
        for cell in row.cells:
            if 'AMZ' in cell.text:
                has_amz = True
                break
    
    if has_amz or 'Seed' in ' '.join(header) or 'TP' in ' '.join(header) or 'Pooled' in str([c.text for r in table.rows for c in r.cells]):
        print(f"  >>> AMZ/SEED/POOLED ICERIGI TESPIT EDILDI - DETAYLI DUMP:")
        for j, row in enumerate(table.rows):
            cells = [cell.text.strip()[:25] for cell in row.cells]
            print(f"    Satir {j}: {' | '.join(cells)}")

# B2 ve B3 olarak adlandirilmis tablolari bul
print("\n\n" + "=" * 80)
print("PARAGRAF TARAMASI — 'Table B2', 'Table B3', 'Tablo' REFERANSLARI")
print("=" * 80)
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and ('Table' in text or 'Tablo' in text or 'B2' in text or 'B3' in text):
        print(f"  P{i}: {text[:150]}")
