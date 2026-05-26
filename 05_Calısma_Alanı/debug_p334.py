# -*- coding: utf-8 -*-
"""Run bazli detayli analiz P334"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from pathlib import Path

fpath = Path(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\BES_Pension_Fund_Report.docx")
doc = Document(str(fpath))

p = doc.paragraphs[334]
print(f"P334 run sayisi: {len(p.runs)}")
for i, run in enumerate(p.runs):
    print(f"  Run {i}: font={run.font.name}, size={run.font.size}, bold={run.bold}")
    print(f"    Text: '{run.text}'")

# "18/(18+3)" veya "85.71" iceren tum paragraflar
print("\n--- 18/(18+3) veya 85.71 iceren paragraflar ---")
for i, p in enumerate(doc.paragraphs):
    if "18/(18+3)" in p.text or ("85.71" in p.text and "Spec" in p.text):
        print(f"  P{i}: {p.text[:200]}")
