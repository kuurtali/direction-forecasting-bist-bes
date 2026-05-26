# -*- coding: utf-8 -*-
"""
HATA DÜZELTME — Pooled satırı Acc=80.21% (77/96) olarak geri al
CM zaten TP=60, TN=17 (=77) ile uyumlu — sadece Acc sütununu düzelt
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from pathlib import Path

fpath = Path(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\BES_Pension_Fund_Report.docx")
doc = Document(str(fpath))

# Tablo B3 = Tablo 23 (index 22)
table_b3 = doc.tables[22]
pooled_row = table_b3.rows[5]

# Mevcut durumu göster
print("MEVCUT POOLED SATIRI:")
for i, cell in enumerate(pooled_row.cells):
    print(f"  Sutun {i}: '{cell.text.strip()}'")

# Sutun 1 (Acc): "84.38% (81/96)" -> "80.21% (77/96)"
acc_cell = pooled_row.cells[1]
para = acc_cell.paragraphs[0]
if para.runs:
    para.runs[0].text = "80.21% (77/96)"
    for r in para.runs[1:]:
        r.text = ""
print("\nDÜZELTME: Acc -> 80.21% (77/96)")

# Sutun 2 (Sens): "80.00% (60/75)" — zaten doğru (true pooled)
# Sutun 3 (Spec): "80.95% (17/21)" — zaten doğru (true pooled)
# Sutun 4 (CM): "TP=60, FN=15, TN=17, FP=4" — zaten doğru (sums to 77, matches 80.21%)

doc.save(str(fpath))

# P334 düzelt — "Pooled CM diagonal = 81/96 = 84.38%" kısmını kaldır
doc = Document(str(fpath))
p = doc.paragraphs[334]
for run in p.runs:
    if "Pooled CM diagonal = 81/96 = 84.38%." in run.text:
        run.text = run.text.replace(
            "Pooled Accuracy = 77/96 = 80.21% (3-seed mean); Pooled CM diagonal = 81/96 = 84.38%.",
            "Pooled Accuracy = 77/96 = 80.21% (sum of per-seed correct predictions: 27+23+27 = 77)."
        )
        print(f"P334 DUZELTILDI")
doc.save(str(fpath))

# SON DOGRULAMA
print("\n" + "=" * 80)
print("SON DOGRULAMA")
print("=" * 80)
doc = Document(str(fpath))
t = doc.tables[22]
for i, row in enumerate(t.rows):
    cells = [c.text.strip()[:35] for c in row.cells]
    tag = ""
    if i >= 1 and i <= 3:
        cm = cells[4]
        nums = [int(x.split("=")[1]) for x in cm.replace(",","").split() if "=" in x]
        total = sum(nums)
        correct = nums[0] + nums[2] if len(nums) >= 3 else 0
        tag = f" [N={total} {'OK' if total==32 else 'HATA'}] [Dogru={correct}]"
    elif i == 5:
        cm = cells[4]
        nums = [int(x.split("=")[1]) for x in cm.replace(",","").split() if "=" in x]
        total = sum(nums)
        correct = nums[0] + nums[2] if len(nums) >= 3 else 0
        tag = f" [N={total} {'OK' if total==96 else 'HATA'}] [Dogru={correct} -> {correct}/96={correct/96*100:.2f}%]"
    print(f"  Satir {i}: {' | '.join(cells)}{tag}")

print(f"\nP334: {doc.paragraphs[334].text[:250]}")
