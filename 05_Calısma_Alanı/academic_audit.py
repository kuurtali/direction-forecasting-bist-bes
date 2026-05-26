# -*- coding: utf-8 -*-
"""MAKALE.docx — Akademik Yayın Hazırlık Denetimi (Eleştirel)"""
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx")

issues = []
warnings = []
info = []

# ═══════════════════════════════════════════════════════════════
# 1. YAPI — IMRaD kontrolü
# ═══════════════════════════════════════════════════════════════
print("=" * 70)
print("1. YAPI VE BÖLÜM KONTROLÜ")
print("=" * 70)

sections_found = {}
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    for n in range(1, 7):
        if t.startswith(str(n) + ". ") and len(t) < 100:
            is_sub = any(t.startswith(str(n) + "." + str(m)) for m in range(1, 10))
            if not is_sub:
                sections_found[n] = (i, t)
                print(f"  Bolum {n}: P{i} [{t}]")

required = {1: "Giris", 2: "Yontem", 3: "Veri", 4: "Bulgular", 5: "Tartisma", 6: "Sonuc"}
for n, name in required.items():
    if n not in sections_found:
        issues.append(f"Bolum {n} ({name}) bulunamadi!")

# ═══════════════════════════════════════════════════════════════
# 2. ÖZ / ABSTRACT
# ═══════════════════════════════════════════════════════════════
print("\n" + "=" * 70)
print("2. ÖZ / ABSTRACT")
print("=" * 70)

oz_text = ""
abstract_text = ""
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "Öz":
        oz_text = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
    if t == "Abstract":
        # Next non-empty paragraph after Abstract title
        for j in range(i + 1, min(i + 3, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip().startswith("Directional"):
                pass  # English title
            elif len(doc.paragraphs[j].text.strip()) > 50:
                abstract_text = doc.paragraphs[j].text.strip()
                break

oz_words = len(oz_text.split()) if oz_text else 0
abs_words = len(abstract_text.split()) if abstract_text else 0
print(f"  Oz: {oz_words} kelime")
print(f"  Abstract: {abs_words} kelime")

if oz_words < 100:
    warnings.append(f"Oz cok kisa ({oz_words} kelime, min 150 onerisi)")
if oz_words > 300:
    warnings.append(f"Oz cok uzun ({oz_words} kelime, max 250 onerisi)")
if abs_words < 100:
    warnings.append(f"Abstract cok kisa ({abs_words} kelime)")

# ═══════════════════════════════════════════════════════════════
# 3. KAYNAKLAR
# ═══════════════════════════════════════════════════════════════
print("\n" + "=" * 70)
print("3. KAYNAKLAR")
print("=" * 70)

refs = []
ref_nums_in_text = set()
for p in doc.paragraphs:
    t = p.text.strip()
    if re.match(r'^\[\d+\]', t):
        num = int(re.match(r'^\[(\d+)\]', t).group(1))
        refs.append(num)
    # Find in-text citations
    for m in re.finditer(r'\[(\d+)\]', p.text):
        ref_nums_in_text.add(int(m.group(1)))

print(f"  Kaynak sayisi: {len(refs)}")
print(f"  Metinde atif yapilan: {sorted(ref_nums_in_text)}")

# Check sequential
if refs:
    expected = list(range(1, max(refs) + 1))
    missing = set(expected) - set(refs)
    if missing:
        issues.append(f"Eksik kaynak numaralari: {missing}")
    
    # Unreferenced
    unreferenced = set(refs) - ref_nums_in_text
    if unreferenced:
        warnings.append(f"Metinde atif yapilmayan kaynaklar: {sorted(unreferenced)}")
    
    # Cited but not in refs
    uncited = ref_nums_in_text - set(refs)
    if uncited:
        issues.append(f"Kaynak listesinde olmayan atiflar: {sorted(uncited)}")

print(f"  Ardisik mi: {refs == list(range(1, len(refs)+1))}")

# ═══════════════════════════════════════════════════════════════
# 4. ÇİZELGE ve ŞEKİL NUMARALAMA
# ═══════════════════════════════════════════════════════════════
print("\n" + "=" * 70)
print("4. ÇİZELGE ve ŞEKİL NUMARALAMA")
print("=" * 70)

tables_found = []
figures_found = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m_tab = re.match(r'^Çizelge (\d+)\.', t)
    m_fig = re.match(r'^Şekil (\d+)\.', t)
    if m_tab:
        tables_found.append(int(m_tab.group(1)))
    if m_fig:
        figures_found.append(int(m_fig.group(1)))

print(f"  Cizelge numaralari: {tables_found}")
print(f"  Sekil numaralari: {figures_found}")

# Check sequential
if tables_found and tables_found != list(range(1, len(tables_found) + 1)):
    issues.append(f"Cizelge numaralari ardisik degil: {tables_found}")
if figures_found and figures_found != list(range(1, len(figures_found) + 1)):
    issues.append(f"Sekil numaralari ardisik degil: {figures_found}")

# Check in-text references to tables and figures
tab_refs = set()
fig_refs = set()
for p in doc.paragraphs:
    for m in re.finditer(r'Çizelge (\d+)', p.text):
        tab_refs.add(int(m.group(1)))
    for m in re.finditer(r'Şekil (\d+)', p.text):
        fig_refs.add(int(m.group(1)))

# Tables/figures not referenced in text
tab_unreferenced = set(tables_found) - tab_refs
fig_unreferenced = set(figures_found) - fig_refs
if tab_unreferenced:
    warnings.append(f"Metinde referans verilmeyen cizelgeler: {sorted(tab_unreferenced)}")
if fig_unreferenced:
    warnings.append(f"Metinde referans verilmeyen sekiller: {sorted(fig_unreferenced)}")

# ═══════════════════════════════════════════════════════════════
# 5. İÇERİK KALİTE KONTROL
# ═══════════════════════════════════════════════════════════════
print("\n" + "=" * 70)
print("5. İÇERİK KALİTE")
print("=" * 70)

full_text = " ".join(p.text for p in doc.paragraphs)
word_count = len(full_text.split())
print(f"  Toplam kelime: {word_count}")

# Check for p-values mentioned
p_vals = re.findall(r'p\s*[=<]\s*[\d,\.]+', full_text)
print(f"  p-degeri referanslari: {len(p_vals)}")

# Check for common academic problems
if "%" not in full_text:
    issues.append("Yuzde isareti (%) bulunamadi")

# Check for methodology keywords
method_keywords = ["confusion matrix", "accuracy", "sensitivity", "specificity", "balanced accuracy", "naive baseline"]
for kw in method_keywords:
    if kw.lower() not in full_text.lower():
        warnings.append(f"Metodoloji anahtar kelimesi bulunamadi: {kw}")

# English abstract check
if abstract_text:
    # Check for Turkish characters in English abstract
    turkish_chars = set("çğıöşüÇĞİÖŞÜ")
    has_turkish = any(c in abstract_text for c in turkish_chars)
    if has_turkish:
        warnings.append("Ingilizce abstract'ta Turkce karakter var")

# ═══════════════════════════════════════════════════════════════
# 6. MAKALE BAŞLIĞI
# ═══════════════════════════════════════════════════════════════
print("\n" + "=" * 70)
print("6. BAŞLIK")
print("=" * 70)

title = doc.paragraphs[0].text.strip()
title_words = len(title.split())
print(f"  Baslik: {title}")
print(f"  Kelime sayisi: {title_words}")
# Template says max 11 words
if title_words > 11:
    warnings.append(f"Baslik {title_words} kelime — sablon max 11 kelime diyor")

# Template: only first word capitalized
words = title.split()
if len(words) > 1:
    upper_words = [w for w in words[1:] if w[0].isupper() and w not in ["ARIMA", "LSTM", "CNN", "1D-CNN", "BIST", "BES", "THYAO", "Türkiye"]]
    if upper_words:
        warnings.append(f"Baslik: sablon 'yalnizca ilk sozcuk buyuk' diyor, buyuk baslayanlar: {upper_words}")

# ═══════════════════════════════════════════════════════════════
# 7. ORCID / YAZAR BİLGİLERİ
# ═══════════════════════════════════════════════════════════════
print("\n" + "=" * 70)
print("7. YAZAR BİLGİLERİ")
print("=" * 70)

orcid_count = full_text.count("ORCID")
print(f"  ORCID sayisi: {orcid_count}")
if orcid_count < 2:
    warnings.append("ORCID bilgileri eksik olabilir")

# Email
email_count = full_text.count("@")
print(f"  Email sayisi: {email_count}")

# ═══════════════════════════════════════════════════════════════
# SONUÇ
# ═══════════════════════════════════════════════════════════════
print("\n" + "=" * 70)
print("SONUÇ RAPORU")
print("=" * 70)

if issues:
    print(f"\n❌ KRİTİK SORUNLAR ({len(issues)}):")
    for issue in issues:
        print(f"   {issue}")

if warnings:
    print(f"\n⚠️ UYARILAR ({len(warnings)}):")
    for w in warnings:
        print(f"   {w}")

if not issues and not warnings:
    print("\n✅ SORUN YOK — Makale akademik yayin icin hazir.")
elif not issues:
    print(f"\n📋 Kritik sorun yok, {len(warnings)} uyari mevcut.")
