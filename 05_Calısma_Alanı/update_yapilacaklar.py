# -*- coding: utf-8 -*-
"""YAPILACAKLAR_RAPOR.docx güncelleme — tamamlanan maddeleri işaretle ve 2. Seans özeti ekle"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import shutil
from datetime import datetime

fpath = Path(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\YAPILACAKLAR_RAPOR.docx")

# Yedek al
backup = fpath.with_suffix('.docx.backup')
shutil.copy2(str(fpath), str(backup))
print(f"Yedek alindi: {backup}")

doc = Document(str(fpath))

# ─── 1. Mevcut maddeleri durum ile güncelle ───
status_updates = {
    "[H1]": "✅ TAMAMLANDI — Şekil 21 (seed sağlamlık) ve Şekil 23 (BA vs Acc) farklı görsellerdir, kaldırma gerekmez.",
    "[H2]": "⏳ BEKLİYOR — MAKALE.docx Çizelge 18 altına dipnot eklenmeli.",
    "[H3]": "⏳ BEKLİYOR — Anahtar sözcük başharflerinin büyültülmesi.",
    "[H4]": "⏳ BEKLİYOR — Tartışma bölümü alt başlıkları (5.1-5.5).",
    "[D3]": "✅ TAMAMLANDI — BES_Pension_Fund_Report.docx Tablo B2 (4 etiket düzeltmesi) ve B3 (per-seed CM + Pooled satırı) düzeltildi.",
}

updated_count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for tag, new_status in status_updates.items():
        if tag in text:
            # Zaten işaretlenmişse atla
            if "TAMAMLANDI" in p.text or "BEKLİYOR" in p.text:
                print(f"  P{i} [{tag}]: Zaten isaretli, atlaniyor")
                break
            # Son run'a durum ekle
            if p.runs:
                last_run = p.runs[-1]
                last_run.text = last_run.text.rstrip() + f"\n    → {new_status}"
                updated_count += 1
                print(f"  P{i} [{tag}]: Guncellendi")
            break

# ─── 2. Yeni bölüm: 2. Seans düzeltmeleri ───
# Başlık paragrafı
section_title = doc.add_paragraph()
run = section_title.add_run("7. 2. Seans Düzeltmeleri (01.05.2026)")
run.bold = True
run.font.size = Pt(11)

items = [
    "Tablo B3 Düzeltme: Per-seed CM değerleri (TP=27→21, N=38→32), Pooled Acc=80.21% (77/96) tutarlı hale getirildi.",
    "Tablo B2 Düzeltme: 4 satırda Feature_Set etiketi CSV ground-truth ile eşleştirildi (full→technical/closing).",
    "P334 Açıklama: Pooled satırı metni güncellendi — Sens=60/75=80.00%, Spec=17/21=80.95%.",
    "UYIK_2026_BES.pptx: Tarandı, HATA YOK — kongre için hazır.",
    "Şekil 23: Kaldırılmadı — Şekil 21 (seed) ve Şekil 23 (BA vs Acc) farklı görsellerdir.",
    "PROJECT_REPORT.txt: 3281 satır otomatik denetim — 12/12 kritik değer CSV uyumlu, encoding düzeltmeleri yapıldı.",
]

for item in items:
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet'] if 'List Bullet' in [s.name for s in doc.styles] else doc.styles['Normal']
    p.clear()
    run = p.add_run(f"• {item}")
    run.font.size = Pt(10)

# ─── 3. Doğrulama notu ───
note = doc.add_paragraph()
run = note.add_run(f"\n📋 Aritmetik doğrulama (02.05.2026 08:07): 80.21% = Pooled 77/96 (3-seed mean ile aynı), 84.38% = Seed 23/98 bireysel doğruluk (27/32). Her iki değer doğrudur, çelişki yoktur.")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.save(str(fpath))
print(f"\nYAPILACAKLAR_RAPOR.docx guncellendi:")
print(f"  {updated_count} madde isaretlendi")
print(f"  {len(items)} yeni 2. Seans maddesi eklendi")
print(f"  1 dogrulama notu eklendi")
