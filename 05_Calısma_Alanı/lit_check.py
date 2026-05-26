# -*- coding: utf-8 -*-
"""02 klasörü vs MAKALE kaynakları — tam eşleştirme"""
import sys, io, glob
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

# ═══ 1. Literatür taraması PDF içerik tarama ═══
print("=== 02 KLASÖRÜNDEKİ PDF'LER ===")
for pdf in glob.glob(r"c:\Users\Kurt\Desktop\Proje\02_Akademik_Kanıtlar\*.pdf"):
    with open(pdf, "rb") as f:
        raw = f.read()
    fname = pdf.split("\\")[-1]
    print(f"\n{fname} ({len(raw)} bytes):")
    
    authors = [
        "Valverde", "Box", "Jenkins", "Engle", "Bollerslev",
        "Kim", "Khaidem", "Hochreiter", "Fischer", "Selvin",
        "Van der Burgt", "Ozbayoglu", "Yildirim", "Fama",
        "Hamilton", "Chollet", "LeCun", "Bai", "Murphy",
        "Sonmez", "Wang", "Albeladi", "Huang", "Koong",
        "NASDAQ", "ARIMA", "LSTM", "CNN", "Tilburg",
        "BIST", "pension", "emeklilik",
    ]
    found = []
    for a in authors:
        if a.lower().encode() in raw.lower():
            found.append(a)
    print(f"  Bulunan: {found}")

# ═══ 2. tez_turkce_ceviri.docx — kaynak listesi ═══
print("\n\n=== tez_turkce_ceviri.docx KAYNAKLARI ===")
doc = Document(r"c:\Users\Kurt\Desktop\Proje\02_Akademik_Kanıtlar\tez_turkce_ceviri.docx")
tez_refs = []
for p in doc.paragraphs:
    t = p.text.strip()
    # Look for reference-like entries
    if t and (t[0].isupper() or t.startswith("(")) and ("202" in t or "201" in t or "199" in t or "198" in t):
        if len(t) > 30 and len(t) < 500:
            tez_refs.append(t)
            
# Also check last 50 paragraphs for bibliography
for p in doc.paragraphs[-80:]:
    t = p.text.strip()
    if t and len(t) > 20:
        if any(yr in t for yr in ["2023", "2022", "2021", "2020", "2019", "2018", "2017", "2016", "1997", "1994", "1970"]):
            if t not in tez_refs:
                tez_refs.append(t)

for ref in tez_refs:
    print(f"  {ref[:150]}")
print(f"\nToplam tez kaynak adayı: {len(tez_refs)}")

# ═══ 3. DOC1 Terimler ═══
print("\n\n=== DOC1 Terimler Formuller.docx ===")
doc1 = Document(r"c:\Users\Kurt\Desktop\Proje\02_Akademik_Kanıtlar\DOC1 Terimler Formuller.docx")
for i in range(min(15, len(doc1.paragraphs))):
    t = doc1.paragraphs[i].text.strip()
    if t:
        print(f"  P{i}: {t[:150]}")

# ═══ 4. Makale_Ornegi PDF — kaynak karşılaştırma ═══
print("\n\n=== Makale_Ornegi (05 klasörü) ===")
ornek_pdfs = glob.glob(r"c:\Users\Kurt\Desktop\Proje\05_Hoca_Onerlileri_Ve_Ornekler\*.pdf")
for pdf in ornek_pdfs:
    with open(pdf, "rb") as f:
        raw = f.read()
    fname = pdf.split("\\")[-1][:60]
    print(f"{fname} ({len(raw)} bytes)")

# ═══ 5. abstract_uyik docx ═══
print("\n\n=== abstract_uyik---2026.docx ===")
try:
    abst = Document(r"c:\Users\Kurt\Desktop\Proje\05_Hoca_Onerlileri_Ve_Ornekler\abstract_uyik---2026_26.04.26_ÖKE.docx")
    for i in range(min(20, len(abst.paragraphs))):
        t = abst.paragraphs[i].text.strip()
        if t:
            print(f"  P{i}: {t[:200]}")
except Exception as e:
    print(f"  Hata: {e}")
