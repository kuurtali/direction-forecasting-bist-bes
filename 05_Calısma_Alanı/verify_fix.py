# -*- coding: utf-8 -*-
"""DÜZELTME DOĞRULAMA — B2 + B3 sonrası kontrol"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from pathlib import Path

doc = Document(str(Path(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\BES_Pension_Fund_Report.docx")))

print("=" * 80)
print("DOGRULAMA: TABLO B3 (Tablo 23)")
print("=" * 80)
t = doc.tables[22]
for i, row in enumerate(t.rows):
    cells = [c.text.strip()[:30] for c in row.cells]
    tag = ""
    if i >= 1 and i <= 3:
        # Per-seed: CM toplami = 32 olmali
        cm = cells[4]
        nums = [int(x.split("=")[1]) for x in cm.replace(",", "").split() if "=" in x]
        total = sum(nums) if nums else 0
        tp_tn = nums[0] + nums[2] if len(nums) >= 3 else 0
        tag = f" [N={total}{'  OK' if total==32 else '  HATA!'}] [TP+TN={tp_tn}]"
    elif i == 5:
        cm = cells[4]
        nums = [int(x.split("=")[1]) for x in cm.replace(",", "").split() if "=" in x]
        total = sum(nums) if nums else 0
        tp_tn = nums[0] + nums[2] if len(nums) >= 3 else 0
        tag = f" [N={total}{'  OK' if total==96 else '  HATA!'}] [TP+TN={tp_tn}, Acc={tp_tn}/{total}={tp_tn/total*100:.2f}%]"
    print(f"  Satir {i}: {' | '.join(cells)}{tag}")

print(f"\n{'='*80}")
print("DOGRULAMA: TABLO B2 (Tablo 20 — AMZ)")
print("=" * 80)
t = doc.tables[19]
for i, row in enumerate(t.rows):
    cells = [c.text.strip()[:25] for c in row.cells]
    print(f"  Satir {i}: {' | '.join(cells)}")

# Paragraf 334 kontrol
print(f"\n{'='*80}")
print("DOGRULAMA: P334 Metin")
print("=" * 80)
p = doc.paragraphs[334]
print(f"  {p.text[:200]}")
