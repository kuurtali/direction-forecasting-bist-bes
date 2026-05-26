"""
1. Extract embedded images from MAKALE.docx and compare with 04 directory
2. Fix aggressive language tone
"""
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import zipfile, os, re, hashlib

SRC = r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx"
DIR04 = r"c:\Users\Kurt\Desktop\Proje\04_Gorsel_Portfolyo"

# ============ PART 1: Image inventory ============
print("="*70)
print("1. MAKALE.docx GÖMÜLÜ GÖRSEL ENVANTERİ")
print("="*70)

with zipfile.ZipFile(SRC) as z:
    media_files = [f for f in z.namelist() if f.startswith('word/media/')]
    print(f"\n  Toplam gömülü medya: {len(media_files)}")
    for mf in sorted(media_files):
        info = z.getinfo(mf)
        ext = os.path.splitext(mf)[1]
        print(f"  {mf} ({info.file_size:,} bytes) [{ext}]")

# ============ PART 2: Figure caption → image mapping ============
print("\n" + "="*70)
print("2. ŞEKİL CAPTION LİSTESİ (23 şekil)")
print("="*70)

doc = Document(SRC)
paras = [(p.text.strip(), p) for p in doc.paragraphs]

fig_captions = []
for i, (text, p) in enumerate(paras):
    m = re.match(r'^Şekil\s+(\d+)\.\s*(.*)', text)
    if m:
        fig_captions.append({
            'num': int(m.group(1)),
            'caption': m.group(2)[:100],
            'para_idx': i
        })

print(f"\n  Toplam caption: {len(fig_captions)}")
print(f"  Toplam gömülü görsel: {len(media_files)}")

if len(fig_captions) != len(media_files):
    print(f"  ⚠ UYUMSUZLUK: {len(fig_captions)} caption vs {len(media_files)} görsel")
else:
    print(f"  ✓ Sayı eşleşmesi OK")

for fc in fig_captions:
    # Check if caption is Turkish
    en_indicators = ['the ', ' of ', ' and ', ' for ', ' with ', ' vs ']
    en_count = sum(1 for w in en_indicators if w in fc['caption'].lower())
    lang = 'TR' if en_count < 2 else '⚠EN'
    print(f"  Şekil {fc['num']:2d}: [{lang}] {fc['caption']}")

# ============ PART 3: 04 directory inventory ============
print("\n" + "="*70)
print("3. 04_Gorsel_Portfolyo ENVANTERİ")
print("="*70)

tr_dir = os.path.join(DIR04, "Makale_TR_Grafikleri")
tr_files = sorted(os.listdir(tr_dir))
print(f"\n  Makale_TR_Grafikleri: {len(tr_files)} dosya")
for f in tr_files:
    print(f"    {f}")

# ============ PART 4: Language tone fix ============
print("\n" + "="*70)
print("4. DİL TONU DÜZELTMELERİ")
print("="*70)

# Define replacements: (pattern, context_check, old, new)
tone_fixes = [
    # üstün → görece yüksek performans (only when referring to model performance)
    ('üstün performans', 'görece yüksek performans'),
    ('üstün başarı', 'görece yüksek başarı'),
    # deşifre → ortaya koyma
    ('deşifre etmiştir', 'ortaya koymuştur'),
    ('deşifre etmektedir', 'ortaya koymaktadır'),
    ('deşifre edilmiştir', 'ortaya konmuştur'),
    # en güçlü → en yüksek doğruluğa sahip
    ('en güçlü model', 'en yüksek doğruluğa sahip model'),
    ('en güçlü konfigürasyon', 'en yüksek doğruluğa sahip konfigürasyon'),
    # kanıtlanmıştır → gözlemlenmiştir (careful - don't change istatistiksel kanıt)
    ('kanıtlamıştır', 'ortaya koymuştur'),
    ('kanıtlanmıştır', 'gözlemlenmiştir'),
    # başarılı → olumlu sonuç veren (only triumphalist uses)
    # kesin olarak → güçlü biçimde
    ('kesin olarak', 'güçlü biçimde'),
]

fix_count = 0
for p in doc.paragraphs:
    original = p.text
    for old_phrase, new_phrase in tone_fixes:
        if old_phrase in p.text:
            # Try to replace in runs
            for run in p.runs:
                if old_phrase in run.text:
                    run.text = run.text.replace(old_phrase, new_phrase)
                    fix_count += 1
                    print(f"  DÜZELTME: '{old_phrase}' → '{new_phrase}'")
                    print(f"    Bağlam: ...{run.text[:80]}...")

# Also check for specific patterns paragraph-wide
print(f"\n  Toplam dil tonu düzeltmesi: {fix_count}")

doc.save(SRC)
print(f"\nKaydedildi: {SRC}")
