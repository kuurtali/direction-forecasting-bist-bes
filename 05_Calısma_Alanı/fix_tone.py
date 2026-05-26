"""Language tone softening - paragraph-level text replacement"""
from docx import Document
import re

SRC = r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx"
doc = Document(SRC)

# First, find ALL instances of aggressive language with full context
print("MEVCUT AGRESIF IFADELER (tam bağlam):")
print("="*70)

patterns = [
    (r'üstün\w*', 'ustun'),
    (r'deşifre\w*', 'desifre'),
    (r'kanıtla\w+', 'kanitla'),
    (r'en güçlü', 'en_guclu'),
    (r'kesin olarak', 'kesin_olarak'),
    (r'kesinlikle', 'kesinlikle'),
    (r'başarılı\b', 'basarili'),
    (r'ispat\w+', 'ispat'),
    (r'tartışmasız', 'tartismasiz'),
    (r'şüphesiz', 'suphesiz'),
    (r'çarpıcı', 'carpici'),
    (r'dramatik', 'dramatik'),
    (r'olağanüstü', 'olaganüstü'),
]

all_hits = []
for i, p in enumerate(doc.paragraphs):
    text = p.text
    if not text.strip():
        continue
    for pattern, label in patterns:
        for m in re.finditer(pattern, text, re.IGNORECASE):
            # Get context around match
            start = max(0, m.start()-40)
            end = min(len(text), m.end()+40)
            context = text[start:end]
            all_hits.append((i, label, m.group(), context))
            
print(f"Toplam bulgu: {len(all_hits)}\n")
for idx, label, word, ctx in all_hits:
    print(f"  P{idx} [{label}]: ...{ctx}...")
    print()

# Now apply fixes paragraph by paragraph
print("\n" + "="*70)
print("DÜZELTMELERİ UYGULAMA")
print("="*70)

fixes_applied = []

for p in doc.paragraphs:
    original = p.text
    
    # Skip bibliography and section headings
    if re.match(r'^\[\d+\]', original) or re.match(r'^\d+\.\d*\s', original):
        continue
    if not original.strip():
        continue
    
    new_text = original
    
    # Specific replacements with context awareness
    replacements = [
        # "üstün performans" → "yüksek performans"
        ('üstün performans', 'yüksek performans'),
        ('üstün doğruluk', 'yüksek doğruluk'),
        ('üstünlük', 'avantaj'),
        # "deşifre" → "ortaya koy"
        ('deşifre etmiştir', 'ortaya koymuştur'),
        ('deşifre etmektedir', 'ortaya koymaktadır'),
        ('deşifre edilmiştir', 'ortaya konmuştur'),
        ('deşifre etti', 'ortaya koydu'),
        # "kanıtla" → "göster/ortaya koy" (NOT istatistiksel context)
        ('kanıtlamıştır', 'ortaya koymuştur'),
        ('kanıtlanmıştır', 'gözlemlenmiştir'),
        ('kanıtlamaktadır', 'göstermektedir'),
        # "en güçlü" → "en yüksek doğruluğa sahip"
        ('en güçlü model', 'en yüksek doğruluğa sahip model'),
        ('en güçlü konfigürasyon', 'en yüksek doğruluğa sahip konfigürasyon'),
        ('en güçlü adayı', 'en yüksek doğruluk değerine sahip adayı'),
        # kesin olarak → güçlü biçimde
        ('kesin olarak', 'güçlü biçimde'),
    ]
    
    for old, new in replacements:
        if old in new_text:
            new_text = new_text.replace(old, new)
            fixes_applied.append(f"'{old}' → '{new}'")
    
    # Apply if changed
    if new_text != original:
        # Rebuild paragraph runs
        if p.runs:
            # Simple approach: put all text in first run
            first_run_format = p.runs[0]
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''

# Save
doc.save(SRC)

print(f"\n  Uygulanan düzeltmeler: {len(fixes_applied)}")
for f in fixes_applied:
    print(f"    ✓ {f}")

print(f"\nKaydedildi: {SRC}")
