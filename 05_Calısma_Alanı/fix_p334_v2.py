# -*- coding: utf-8 -*-
"""P334 Specificity kalan düzeltme"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from pathlib import Path

fpath = Path(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\BES_Pension_Fund_Report.docx")
doc = Document(str(fpath))

p = doc.paragraphs[334]
for run in p.runs:
    if "18/(18+3)" in run.text or "85.71%" in run.text:
        old = run.text
        run.text = run.text.replace("Specificity = 18/(18+3) = 85.71%", 
                                     "Pooled Sens = 60/75 = 80.00%, Pooled Spec = 17/21 = 80.95%")
        if old != run.text:
            print(f"FIXED: '{old}' -> '{run.text}'")

doc.save(str(fpath))

# Final check
doc2 = Document(str(fpath))
print(f"\nFINAL P334: {doc2.paragraphs[334].text[:350]}")
