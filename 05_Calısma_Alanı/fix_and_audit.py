"""
Fix 1: Abstract italic
Fix 2: Keywords alphabetical order (TR + EN)
Audit: All figures check + language tone analysis
"""
from docx import Document
from docx.shared import Pt
import re, shutil
from datetime import datetime

SRC = r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx"
ts = datetime.now().strftime("%Y%m%d_%H%M%S")
shutil.copy2(SRC, SRC.replace("MAKALE.docx", f"MAKALE_yedek_{ts}.docx"))

doc = Document(SRC)

# ============ FIX 1: Make English Abstract italic ============
print("--- FIX 1: Abstract italic ---")
abstract_found = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Abstract':
        abstract_found = True
        continue
    if abstract_found and p.text.strip():
        # This is the English abstract paragraph - make all runs italic
        if p.text.strip().startswith('Keywords:') or p.text.strip().startswith('Anahtar'):
            break
        for run in p.runs:
            run.italic = True
        print(f"  Made italic: {p.text[:80]}...")
        abstract_found = True  # Continue for next paragraphs of abstract
    if abstract_found and p.text.strip().startswith('Keywords:'):
        break

print("  Abstract italic DONE")

# ============ FIX 2: Keywords alphabetical ============
print("\n--- FIX 2: Keywords alphabetical ---")

for p in doc.paragraphs:
    t = p.text.strip()
    
    # Turkish keywords
    if t.startswith('Anahtar'):
        prefix = t.split(':')[0] + ': '
        kws_text = t.split(':', 1)[1].strip().rstrip('.')
        kws = [k.strip() for k in kws_text.split(',')]
        
        # Turkish alphabetical sort using proper TDK order
        def tr_sort_key(s):
            order = 'AaBbCcÇçDdEeFfGgĞğHhIıİiJjKkLlMmNnOoÖöPpRrSsŞşTtUuÜüVvYyZz0123456789'
            return [order.index(c) if c in order else 999 for c in s]
        sorted_kws = sorted(kws, key=tr_sort_key)
        
        new_text = prefix + ', '.join(sorted_kws) + '.'
        print(f"  TR Eski:  {t[:100]}")
        print(f"  TR Yeni:  {new_text[:100]}")
        
        # Replace
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''
    
    # English keywords
    if t.startswith('Keywords:'):
        prefix = 'Keywords: '
        kws_text = t.split(':', 1)[1].strip().rstrip('.')
        kws = [k.strip() for k in kws_text.split(',')]
        sorted_kws = sorted(kws, key=str.lower)
        
        new_text = prefix + ', '.join(sorted_kws) + '.'
        print(f"  EN Eski:  {t[:100]}")
        print(f"  EN Yeni:  {new_text[:100]}")
        
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''

# ============ SAVE ============
doc.save(SRC)
print(f"\nSaved to {SRC}")

# ============ AUDIT: Figure analysis ============
print("\n" + "="*70)
print("FIGURE ANALYSIS - All 23 figures")
print("="*70)

doc2 = Document(SRC)
paras = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]

for p in paras:
    m = re.match(r'^(Sekil|Şekil)\s+(\d+)\.\s*(.*)', p, re.IGNORECASE)
    if m:
        num = m.group(2)
        caption = m.group(3)[:120]
        # Check language
        en_words = ['the','of','and','for','with','vs','model','accuracy','seed']
        has_english = sum(1 for w in en_words if f' {w} ' in f' {caption.lower()} ')
        lang = 'EN?' if has_english > 2 else 'TR'
        print(f"  Sekil {num}: [{lang}] {caption}")

# ============ AUDIT: Language tone analysis ============
print("\n" + "="*70)
print("LANGUAGE TONE ANALYSIS")
print("="*70)

# Find aggressive/triumphalist language
aggressive_patterns = [
    (r'kanıtla[nmrd]', 'kanitlandi'),
    (r'başar[ıi]l[ıi]', 'basarili'),
    (r'üstün', 'ustun'),
    (r'mükemmel', 'mukemmel'),
    (r'kesin\b', 'kesin'),
    (r'tartışmasız', 'tartismasiz'),
    (r'şüphesiz', 'suphesiz'),
    (r'ispatla[nmrd]', 'ispatlandi'),
    (r'yok ett?i', 'yok etti'),
    (r'ezdi', 'ezdi'),
    (r'olağanüstü', 'olaganüstü'),
    (r'çarpıcı', 'carpici'),
    (r'dramatik', 'dramatik'),
    (r'devasa', 'devasa'),
    (r'muazzam', 'muazzam'),
    (r'yıldız bulgu', 'yildiz bulgu'),
    (r'en güçlü', 'en guclu'),
    (r'deşifre', 'desifre'),
]

tone_issues = []
for i, p_text in enumerate(paras):
    for pattern, label in aggressive_patterns:
        if re.search(pattern, p_text, re.IGNORECASE):
            tone_issues.append((i, label, p_text[:120]))

print(f"\n  Bulunan agresif/zafer tonu ifadeleri: {len(tone_issues)}")
seen_labels = set()
for idx, label, txt in tone_issues:
    if label not in seen_labels:
        seen_labels.add(label)
        print(f"  [{label}] P{idx}: {txt}")

# Academic alternatives suggestion
print("\n  YUMUSATMA ÖNERİLERİ:")
suggestions = {
    'kanitlandi': '"kanıtlanmıştır" → "gözlemlenmiştir / elde edilmiştir / ortaya konmuştur"',
    'basarili': '"başarılı" → "anlamlı performans gösteren / olumlu sonuç veren"',
    'ustun': '"üstün" → "görece yüksek performans sergileyen"',
    'desifre': '"deşifre" → "sistematik biçimde raporlanması"',
    'en guclu': '"en güçlü" → "en yüksek doğruluk değerine sahip"',
    'yildiz bulgu': '"yıldız bulgu" → "merkezi bulgu / temel sonuç"',
    'dramatik': '"dramatik" → "belirgin / kayda değer"',
    'carpici': '"çarpıcı" → "dikkat çekici / kayda değer"',
    'olaganüstü': '"olağanüstü" → "kayda değer düzeyde yüksek"',
}
for label, suggestion in suggestions.items():
    if label in seen_labels:
        print(f"    {suggestion}")
