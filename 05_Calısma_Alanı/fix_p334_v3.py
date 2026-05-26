# -*- coding: utf-8 -*-
"""P334 Specificity düzeltme — 3 run'a bölünmüş metni düzelt"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from pathlib import Path

fpath = Path(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\BES_Pension_Fund_Report.docx")
doc = Document(str(fpath))

p = doc.paragraphs[334]
# Run 0 sonunda: "...Specificity = 18"
# Run 1: "/("
# Run 2: "18+3) = 85.71%."
# Hedef: "...Pooled Sens = 60/75, Pooled Spec = 17/21."

r0 = p.runs[0]
# "Specificity = 18" kısmını değiştir
r0.text = r0.text.replace("Specificity = 18", "Pooled Sens = 60/75 = 80.00%, Pooled Spec = 17/21 = 80.95%.")
p.runs[1].text = ""  # "/(" sil
p.runs[2].text = ""  # "18+3) = 85.71%." sil

doc.save(str(fpath))

# Verify
doc2 = Document(str(fpath))
print(f"FINAL P334: {doc2.paragraphs[334].text}")
