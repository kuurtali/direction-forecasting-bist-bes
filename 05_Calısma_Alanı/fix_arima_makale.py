# -*- coding: utf-8 -*-
"""MAKALE.docx — ARIMA 'dirençli/muaf' 3 noktada düzeltme"""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

MAKALE_PATH = r'c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx'
doc = Document(MAKALE_PATH)

# Nonblank indeks -> gerçek paragraf indeksi eşlemesi
nonblank_map = [(i, p) for i, p in enumerate(doc.paragraphs) if p.text.strip()]

# nonblank[85] = doc.paragraphs[103], [86]=105, [136]=176
p85  = doc.paragraphs[103]
p86  = doc.paragraphs[105]
p136 = doc.paragraphs[176]

print("=== ÖNCE ===")
print(f"[85] son 300: ...{p85.text[-300:]}")
print(f"[86]: {p86.text[:200]}")
print(f"[136] 'dirençli' var mı: {'dirençli' in p136.text}")
print()

# ───────────────────────────────────────────────────────────────────
# DÜZELTME 1: [85] — ARIMA uyarı cümlesi ekle (son run'a append)
# ───────────────────────────────────────────────────────────────────
EK_CUMLE = (
    " Ancak bu tablo dikkatli yorumlanmalıdır: ARIMA her iki dönemde de"
    " rassal tahmin sınırında (~%50) seyretmiş olup başarısız kalmaya"
    " devam etmek kavramsal sürüklenmeden muafiyet sayılamaz."
    " Asıl darbeyi yiyen, 2018–2022 döneminde %55–57 performans gösteren"
    " derin öğrenme modelleridir."
)
if p85.runs:
    p85.runs[-1].text += EK_CUMLE
else:
    p85.add_run(EK_CUMLE)
print("✅ Düzeltme 1 — [85] uyarı cümlesi eklendi")

# ───────────────────────────────────────────────────────────────────
# DÜZELTME 2: [86] Şekil 9 başlığı — "muaf kalırken" → dürüst ifade
# ───────────────────────────────────────────────────────────────────
OLD2 = ("ARIMA sürüklenmeden görece muaf kalırken derin öğrenme modelleri"
        " %5–14 puanlık düşüş yaşamıştır")
NEW2 = ("Derin öğrenme modelleri %5–14 puanlık düşüş yaşarken, ARIMA her"
        " iki dönemde de ~%50 (rassal tahmin sınırı) düzeyinde seyretmiştir"
        " — başarısız kalmaya devam etmek sürüklenmeden muafiyet sayılamaz")

old_text86 = p86.text
if OLD2 in old_text86:
    new_text86 = old_text86.replace(OLD2, NEW2)
    # Tüm run metinlerini temizle, ilk run'a tam metni yaz
    for run in p86.runs:
        run.text = ''
    if p86.runs:
        p86.runs[0].text = new_text86
    else:
        p86.add_run(new_text86)
    print("✅ Düzeltme 2 — [86] Şekil 9 başlığı düzeltildi")
else:
    print(f"⚠️  [86] eşleşme bulunamadı. Metin: {old_text86[:150]}")

# ───────────────────────────────────────────────────────────────────
# DÜZELTME 3: [136] 5.3 tartışma — "dirençli kalması" → dürüst ifade
# ───────────────────────────────────────────────────────────────────
OLD3 = ("ARIMA'nın drift'e dirençli kalması, modelin yapısal basitliğinden"
        " ve fark alma adımının doğal olarak rejim değişikliklerine uyum"
        " sağlamasından kaynaklanır.")
NEW3 = ("ARIMA'nın bu süreçte ~%50 düzeyinde stabil kalması, modelin"
        " yapısal basitliğinden kaynaklanmakla birlikte, bu stabilite bir"
        " sürüklenme direnci olarak değil, her iki dönemde de rassal tahmin"
        " sınırında seyretmenin sürdürülmesi olarak yorumlanmalıdır."
        " Başarısız kalmaya devam etmek metodolojik bir üstünlük değildir.")

old_text136 = p136.text
if OLD3 in old_text136:
    new_text136 = old_text136.replace(OLD3, NEW3)
    for run in p136.runs:
        run.text = ''
    if p136.runs:
        p136.runs[0].text = new_text136
    else:
        p136.add_run(new_text136)
    print("✅ Düzeltme 3 — [136] 'dirençli kalması' → doğru ifade")
else:
    # Kısmi arama
    if "dirençli" in old_text136:
        print(f"⚠️  [136] tam eşleşme yok ama 'dirençli' var. Manuel düzelt.")
        idx = old_text136.find("dirençli")
        print(f"  Bağlam: ...{old_text136[max(0,idx-80):idx+120]}...")
    else:
        print("⚠️  [136] 'dirençli' kelimesi bulunamadı — zaten düzeltilmiş olabilir")

# ─── KAYDET ───
doc.save(MAKALE_PATH)
print()
print("💾 MAKALE.docx kaydedildi.")
print()

# ─── DOĞRULAMA ───
doc2 = Document(MAKALE_PATH)
p85v  = doc2.paragraphs[103]
p86v  = doc2.paragraphs[105]
p136v = doc2.paragraphs[176]
print("=== SONRA ===")
print(f"[85] uyarı var mı: {'muafiyet sayılamaz' in p85v.text}")
print(f"[86] 'muaf' YOK mu: {'muaf kalırken' not in p86v.text}")
print(f"[136] 'dirençli' YOK mu: {'dirençli kalması' not in p136v.text}")
