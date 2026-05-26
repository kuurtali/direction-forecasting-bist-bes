from docx import Document
from docx.shared import Pt
import re

doc = Document(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx")
fixes = []

for p in doc.paragraphs:
    orig = p.text
    new = orig

    # FIX 1: Başlıkta "majority class illüzyonu" → "majority class (çoğunluk sınıfı) illüzyonu"
    if "majority class illüzyonu" in new and new.startswith("Türkiye"):
        new = new.replace("majority class illüzyonu", "majority class (çoğunluk sınıfı) illüzyonu")
        fixes.append("BASLIK: majority class → majority class (çoğunluk sınıfı)")

    # FIX 2: Çizelge 1 BES dönem
    if "2021-2026" in new and ("Nis" not in new) and ("Çizelge 1" in new or "veri dönemi" in new.lower() or "BES" in new):
        new = new.replace("2021-2026", "Nisan 2021 – Nisan 2026 (~262 hafta)")
        fixes.append(f"BES DONEM: 2021-2026 → Nisan 2021 – Nisan 2026 (~262 hafta)")

    # FIX 3: Add ARIMA=Naive note in discussion (§5.3 area)
    # Already partially there at P177 based on earlier audit
    # Let's check and add if the exact note isn't there
    if "ARIMA" in new and "Naive" in new and "%78,79" in new and "momentum" not in new and "bağımsız bir öngörü" in new:
        addition = " Bu eşitlik, ARIMA'nın söz konusu ufukta bağımsız bir sinyal üretmediğini, yalnızca serinin momentum yapısını yansıttığını göstermektedir."
        if "momentum" not in new:
            new = new.rstrip() + addition
            fixes.append("TARTISMA: ARIMA=Naive momentum notu eklendi")

    if new != orig:
        if p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""

doc.save(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx")

print(f"Toplam {len(fixes)} düzeltme:")
for f in fixes:
    print(f"  ✓ {f}")

if not fixes:
    print("  Hiç eşleşme bulunamadı — detaylı arama yapılıyor...")
    # Debug: show relevant paragraphs
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if "majority class" in t.lower():
            print(f"  P{i} [majority]: {t[:100]}")
        if "2021-2026" in t:
            print(f"  P{i} [2021-2026]: {t[:100]}")
        if "78,79" in t and "ARIMA" in t:
            print(f"  P{i} [ARIMA+Naive]: {t[:100]}")
