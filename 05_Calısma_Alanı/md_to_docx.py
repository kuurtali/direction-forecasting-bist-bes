import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Stil ayarları
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for s in ['Heading 1','Heading 2','Heading 3']:
    hs = doc.styles[s]
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0,0,0)

with open(r'C:\Users\Hp\Desktop\Proje\MAKALE_FORMAT_TR.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

def add_table(doc, rows):
    if not rows: return
    headers = rows[0]
    ncols = len(headers)
    table = doc.add_table(rows=len(rows), cols=ncols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            cell_text = cell.strip().replace('**','').replace('★','*').replace('⚠','(!)').replace('↓','v').replace('↑','^')
            p = table.rows[i].cells[j].paragraphs[0]
            p.text = cell_text
            p.style.font.size = Pt(10)
            run = p.runs[0] if p.runs else p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            if i == 0:
                run.bold = True

table_rows = []
in_table = False
skip_separator = False

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n').rstrip('\r')
    
    # Skip empty lines between sections
    if line.strip() == '---':
        i += 1
        continue
    
    # Table detection
    if '|' in line and line.strip().startswith('|'):
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # Skip separator rows
        if cells and all(re.match(r'^[-:]+$', c) for c in cells):
            i += 1
            continue
        table_rows.append(cells)
        # Check if next line is still table
        if i+1 < len(lines) and '|' in lines[i+1]:
            i += 1
            continue
        else:
            add_table(doc, table_rows)
            table_rows = []
            i += 1
            continue
    else:
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []
    
    # Headers
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        p = doc.add_heading(title, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue
    
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue
    
    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue
    
    # Bold paragraph (like table titles)
    if line.startswith('**') and line.endswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(line.replace('**',''))
        run.bold = True
        run.font.size = Pt(11)
        i += 1
        continue
    
    # Blockquote
    if line.startswith('> '):
        text = line[2:].replace('**','').replace('★','*')
        p = doc.add_paragraph(text, style='Intense Quote' if 'Intense Quote' in [s.name for s in doc.styles] else 'Normal')
        p.paragraph_format.left_indent = Inches(0.5)
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        i += 1
        continue
    
    # Italic note
    if line.startswith('*') and not line.startswith('**'):
        text = line.strip('*').strip()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(10)
        i += 1
        continue
    
    # List items
    if line.startswith('- '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
        i += 1
        continue
    
    if re.match(r'^\d+\. ', line):
        text = re.sub(r'^\d+\. ', '', line)
        p = doc.add_paragraph(text, style='List Number')
        i += 1
        continue
    
    # Normal paragraph
    if line.strip():
        clean = line.replace('**','').replace('★','*').replace('⚠','(!)')
        doc.add_paragraph(clean)
    
    i += 1

# Flush remaining table
if table_rows:
    add_table(doc, table_rows)

out = r'C:\Users\Hp\Desktop\Proje\MAKALE_FORMAT_TR.docx'
doc.save(out)
print(f"SAVED: {out}")
