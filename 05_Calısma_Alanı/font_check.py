# -*- coding: utf-8 -*-
"""MAKALE vs ŞABLON — Yazı tipi boyutu birebir karşılaştırma"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx")
tmpl = Document(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\Makale_Formatı.docx")

# ═══ ŞABLON KURALLARI (template'den çıkarılmış) ═══
# Başlık:           18pt
# Öz / Abstract:    10pt (başlık), 9pt (metin)
# Anahtar sözcük:   9pt Bold+Italic
# Bölüm başlığı:    11pt Bold (before 18, after 18)
# Alt bölüm:        11pt Italic (before 18, after 12)
# Gövde metin:      11pt (before 0, after 12)
# Çizelge/Şekil:    11pt Bold
# Kaynak başlık:    11pt Bold (before 30, after 12)
# Kaynak madde:     10pt
# Formül:           11pt

print("=" * 75)
print("ŞABLON KURALLARI vs MAKALE — HER PARAGRAF")
print("=" * 75)

rules = {
    "baslik":     {"expected": 18.0, "label": "Başlık (18pt)"},
    "oz_title":   {"expected": 10.0, "label": "Öz başlığı (10pt)"},
    "oz_text":    {"expected": 9.0,  "label": "Öz metin (9pt)"},
    "kw":         {"expected": 9.0,  "label": "Anahtar sözcük (9pt B+I)"},
    "abs_title":  {"expected": 10.0, "label": "Abstract başlığı (10pt)"},
    "eng_title":  {"expected": 10.0, "label": "İng. başlık (10pt B)"},
    "abs_text":   {"expected": 9.0,  "label": "Abstract metin (9pt I)"},
    "eng_kw":     {"expected": 9.0,  "label": "Keywords (9pt B+I)"},
    "section":    {"expected": 11.0, "label": "Bölüm başlığı (11pt B)"},
    "subsection": {"expected": 11.0, "label": "Alt bölüm (11pt I)"},
    "body":       {"expected": 11.0, "label": "Gövde (11pt)"},
    "table_fig":  {"expected": 11.0, "label": "Çizelge/Şekil (11pt B)"},
    "ref_title":  {"expected": 11.0, "label": "Kaynaklar başlık (11pt B)"},
    "ref_item":   {"expected": 10.0, "label": "Kaynak madde (10pt)"},
}

issues = []

# Check each critical paragraph
def check_size(para_idx, expected, label, check_bold=None, check_italic=None):
    p = doc.paragraphs[para_idx]
    t = p.text.strip()[:60]
    if not p.runs:
        return
    r = p.runs[0]
    actual = r.font.size.pt if r.font.size else None
    bold = r.font.bold
    italic = r.font.italic
    
    size_ok = actual == expected if actual else True  # inherited = ok
    bold_ok = True if check_bold is None else (bold == check_bold)
    italic_ok = True if check_italic is None else (italic == check_italic)
    
    status = "OK" if (size_ok and bold_ok and italic_ok) else "FAIL"
    
    details = []
    if not size_ok:
        details.append(f"size={actual} (beklenen {expected})")
    if not bold_ok:
        details.append(f"bold={bold} (beklenen {check_bold})")
    if not italic_ok:
        details.append(f"italic={italic} (beklenen {check_italic})")
    
    detail_str = " | ".join(details) if details else ""
    
    if status == "FAIL":
        issues.append(f"P{para_idx}: {label} — {detail_str} [{t}]")
        print(f"  FAIL  P{para_idx}: {label} — {detail_str}")
        print(f"         [{t}]")
    else:
        print(f"  OK    P{para_idx}: {label} (size={actual}, B={bold}, I={italic})")

# ─── BAŞLIK ───
print("\n--- BAŞLIK ---")
check_size(0, 18.0, "Makale başlığı")

# ─── ÖZ ───
print("\n--- ÖZ / ABSTRACT ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "Öz":
        check_size(i, 10.0, "Öz başlığı")
    elif t == "Abstract":
        check_size(i, 10.0, "Abstract başlığı")

# Öz metin (P6)
check_size(6, 9.0, "Öz metin")
# Anahtar sözcükler (P7)
check_size(7, 9.0, "Anahtar sözcükler", check_bold=True, check_italic=True)

# İngilizce başlık (P9)
check_size(9, 10.0, "İngilizce makale başlığı")
# Abstract metin (P10)
check_size(10, 9.0, "Abstract metin", check_italic=True)
# Keywords (P11)
check_size(11, 9.0, "Keywords", check_bold=True, check_italic=True)

# ─── BÖLÜM BAŞLIKLARI ───
print("\n--- BÖLÜM BAŞLIKLARI ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    # Main sections
    is_main = False
    for n in range(1, 7):
        prefix = str(n) + ". "
        if t.startswith(prefix) and len(t) < 100:
            is_sub = any(t.startswith(str(n) + "." + str(m)) for m in range(1, 10))
            if not is_sub:
                is_main = True
                check_size(i, 11.0, f"Bölüm başlığı", check_bold=True)
    
    # Sub sections
    for n in range(1, 7):
        for m in range(1, 10):
            prefix = str(n) + "." + str(m) + "."
            if t.startswith(prefix) and len(t) < 120:
                check_size(i, 11.0, f"Alt bölüm", check_italic=True)

# ─── Kaynaklar başlığı ───
print("\n--- KAYNAKLAR ---")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Kaynaklar":
        check_size(i, 11.0, "Kaynaklar başlığı", check_bold=True)
    elif p.text.strip().startswith("[1]"):
        check_size(i, 10.0, "Kaynak [1]")
        break

# ─── ÇİZELGE/ŞEKİL BAŞLIKLARI ───
print("\n--- ÇİZELGE/ŞEKİL BAŞLIKLARI ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if (t.startswith("Çizelge") or t.startswith("Şekil")) and "." in t[:15]:
        if p.runs:
            r = p.runs[0]
            actual = r.font.size.pt if r.font.size else None
            bold = r.font.bold
            if actual and actual != 11.0:
                issues.append(f"P{i}: Cizelge/Sekil boyut={actual} (beklenen 11) [{t[:40]}]")
                print(f"  FAIL  P{i}: size={actual} [{t[:50]}]")
            elif not bold:
                issues.append(f"P{i}: Cizelge/Sekil bold=False [{t[:40]}]")
                print(f"  FAIL  P{i}: bold=False [{t[:50]}]")

# ─── GÖVDE METİN ÖRNEKLEMESİ ───
print("\n--- GÖVDE METİN (örnekleme) ---")
body_wrong = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if len(t) > 150 and p.runs:  # Long paragraphs = body text
        r = p.runs[0]
        actual = r.font.size.pt if r.font.size else None
        if actual and actual not in [11.0, 9.0, 10.0, 8.0]:  # 9=öz/abs, 10=ref/eng, 8=footnote
            body_wrong.append((i, actual, t[:50]))

if body_wrong:
    for pi, fs, txt in body_wrong:
        print(f"  FAIL  P{pi}: size={fs} [{txt}]")
        issues.append(f"P{pi}: Gövde boyut={fs} (beklenen 11) [{txt}]")
else:
    print(f"  OK    Tüm gövde paragrafları uygun")

# ─── DİPNOT/EK PARAGRAFLAR ───
print("\n--- DİPNOT PARAGRAFLARI ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if p.runs:
        r = p.runs[0]
        actual = r.font.size.pt if r.font.size else None
        if actual and actual < 9.0:
            print(f"  INFO  P{i}: size={actual} [{t[:60]}]")
            if actual < 8.0:
                issues.append(f"P{i}: Çok küçük font={actual}pt [{t[:40]}]")

# ═══ SONUÇ ═══
print("\n" + "=" * 75)
print("SONUÇ")
print("=" * 75)
if issues:
    print(f"\n❌ {len(issues)} SORUN:")
    for iss in issues:
        print(f"   {iss}")
else:
    print("\n✅ Tüm yazı tipi boyutları şablonla uyumlu.")
