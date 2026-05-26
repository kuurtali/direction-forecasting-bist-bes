from docx import Document
import re

doc = Document(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx")

# Find §6 Sonuç section, add limitation reminder before Kaynaklar
added = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    # Find last paragraph before Kaynaklar
    if "Kaynaklar" in text or "KAYNAKÇA" in text.upper():
        # Insert limitation sentence before this
        # Go back to find the last real paragraph of §6
        for j in range(i-1, 0, -1):
            prev = doc.paragraphs[j].text.strip()
            if len(prev) > 30:
                # Add after this paragraph
                # We need to add a new paragraph after j
                new_text = ("Son olarak, bu bulguların yorumlanmasında §5.5'te tartışılan "
                           "test seti boyutu (~32–40 hafta), TEFAS veri kısıtı nedeniyle "
                           "kullanılan proxy göstergeler ve üç tohumla sınırlı seed varyansı "
                           "(SD = 0,072) göz önünde tutulmalıdır.")
                # Add to the end of the last paragraph
                existing = doc.paragraphs[j]
                # Add a new run with the text
                run = existing.add_run(" " + new_text)
                run.font.name = "Times New Roman"
                from docx.shared import Pt
                run.font.size = Pt(11)
                print(f"  Eklendi (P{j} sonuna):")
                print(f"  '{new_text[:100]}...'")
                added = True
                break
        break

if added:
    doc.save(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\MAKALE.docx")
    print("\nKaydedildi.")
else:
    print("HATA: Eklenecek yer bulunamadı!")
