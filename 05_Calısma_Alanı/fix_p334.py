# -*- coding: utf-8 -*-
"""P334 metin düzeltmesi — eski Pooled açıklama güncelle"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from pathlib import Path

doc = Document(str(Path(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\BES_Pension_Fund_Report.docx")))

p = doc.paragraphs[334]
full_text = p.text
print("ONCE:")
print(f"  {full_text[:300]}")

# Run bazlı düzeltme
for run in p.runs:
    t = run.text
    # Eski değerleri düzelt
    t = t.replace("Pooled Accuracy = 81/96 = 84.38%, consistent with the 3-seed mean",
                  "Pooled Accuracy = 77/96 = 80.21% (3-seed mean); Pooled CM diagonal = 81/96 = 84.38%")
    t = t.replace("Specificity = 18/(18+3) = 85.71%",
                  "Pooled Sens = 60/75 = 80.00%, Pooled Spec = 17/21 = 80.95%")
    if t != run.text:
        print(f"  DUZELTILDI: '{run.text[:80]}' -> '{t[:80]}'")
        run.text = t

doc.save(str(Path(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\BES_Pension_Fund_Report.docx")))

# Dogrulama
doc2 = Document(str(Path(r"c:\Users\Kurt\Desktop\Proje\01_Savunma_ve_Ana_Metinler\BES_Pension_Fund_Report.docx")))
print("\nSONRA:")
print(f"  {doc2.paragraphs[334].text[:300]}")
